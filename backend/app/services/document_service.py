"""Document parsing service for extracting text from docx, pdf, xlsx files."""
import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class DocumentService:
    """Parse various document formats into plain text for knowledge extraction."""

    async def parse(self, content: bytes, filename: str, mime_type: Optional[str] = None) -> str:
        """Parse document content to plain text based on file extension."""
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        if ext == "docx":
            return self._parse_docx(content)
        elif ext == "pdf":
            text = self._parse_pdf(content)
            if not text.strip():
                logger.warning("[parse] PyPDF2 extracted 0 chars from %s, falling back to OCR", filename)
                text = await self._parse_pdf_ocr(content)
            return text
        elif ext in ("xlsx", "xls"):
            return self._parse_xlsx(content)
        elif ext == "txt":
            return content.decode("utf-8", errors="replace")
        elif ext == "md":
            return content.decode("utf-8", errors="replace")
        else:
            raise ValueError(f"Unsupported file format: .{ext}. Supported: docx, pdf, xlsx, txt, md")

    def _parse_docx(self, content: bytes) -> str:
        """Extract text from Word document."""
        from docx import Document
        doc = Document(io.BytesIO(content))

        texts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading structure
                if para.style and para.style.name and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "").replace("Heading", "1")
                    try:
                        level_num = int(level)
                    except ValueError:
                        level_num = 1
                    texts.append(f"\n{'#' * level_num} {text}\n")
                else:
                    texts.append(text)

        # Also extract from tables
        for table in doc.tables:
            rows_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows_text.append(" | ".join(cells))
            if rows_text:
                texts.append("\n[Table]\n" + "\n".join(rows_text) + "\n[/Table]\n")

        return "\n".join(texts)

    def _parse_pdf(self, content: bytes) -> str:
        """Extract text from PDF."""
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(content))

        texts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                texts.append(f"--- Page {i+1} ---\n{text.strip()}")

        return "\n\n".join(texts)

    async def _parse_pdf_ocr(self, content: bytes) -> str:
        """Fallback: render PDF pages to images and use AI vision to extract text."""
        import fitz  # pymupdf
        from app.services.claude_service import ClaudeService

        doc = fitz.open(stream=content, filetype="pdf")
        claude = ClaudeService()
        texts = []

        # Process up to 20 pages to stay within reasonable limits
        max_pages = min(len(doc), 20)
        logger.info("[OCR] Processing %d/%d pages via AI vision", max_pages, len(doc))

        for i in range(max_pages):
            page = doc[i]
            # Render at 150 DPI — good balance of quality vs size
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("png")

            try:
                result = await claude.extract_from_image(img_bytes, "image/png")
                page_text = result.get("extracted_text", "")
                if page_text and page_text.strip():
                    texts.append(f"--- Page {i+1} ---\n{page_text.strip()}")
                    logger.info("[OCR] Page %d: %d chars extracted", i+1, len(page_text))
                else:
                    logger.warning("[OCR] Page %d: no text extracted", i+1)
            except Exception as e:
                logger.error("[OCR] Page %d failed: %s", i+1, e)
                continue

        doc.close()
        return "\n\n".join(texts)

    def _parse_xlsx(self, content: bytes) -> str:
        """Extract text from Excel spreadsheet."""
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)

        texts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(cell) if cell is not None else "" for cell in row]
                if any(c.strip() for c in cells):
                    sheet_rows.append(" | ".join(cells))
            if sheet_rows:
                texts.append(f"\n## Sheet: {sheet_name}\n" + "\n".join(sheet_rows))

        wb.close()
        return "\n".join(texts)
